"""Pruebas de la generacion dinamica de documentos Word (.docx).

Verifica para cada endpoint /api/documentos/*:
- 200 con mimetype DOCX y archivo valido (python-docx lo abre).
- La plantilla queda rellenada: sin placeholders {{...}} residuales
  y con los datos reales del registro en el texto.
- 404 cuando el registro no existe y 401 sin token.
"""
import io
from datetime import date

import pytest
from docx import Document

from app import db
from app.models import (
    Cambio, RenovacionProgramada, InformeVisitaTecnica, ActuacionLegal,
    Plano, PlanoRevision,
)

MIME_DOCX = ('application/vnd.openxmlformats-officedocument'
             '.wordprocessingml.document')


def _texto_docx(resp):
    """Valida cabeceras + DOCX y devuelve todo el texto del documento."""
    assert resp.status_code == 200, resp.get_json()
    assert resp.mimetype == MIME_DOCX
    assert 'attachment' in resp.headers.get('Content-Disposition', '')
    doc = Document(io.BytesIO(resp.data))  # falla si no es DOCX valido
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.append(celda.text)
    texto = '\n'.join(partes)
    assert '{{' not in texto and '}}' not in texto, \
        f'Placeholders sin rellenar: {texto}'
    assert 'None' not in texto, 'Un campo vacio se imprimio como "None"'
    return texto


@pytest.fixture()
def datos_documentos(app):
    """Registros de todas las areas para generar cada documento."""
    cambio = Cambio(id_cato=70001, id_afi_titular='4444444',
                    id_afi_nuevo='5555555', tipo_cambio='COMPRA-VENTA',
                    codigo_docu='DOC-99', fecha_cambio=date(2026, 5, 10),
                    resol_nro='RES-7', resol_fecha=date(2026, 5, 12))
    renov = RenovacionProgramada(
        id_renov=900, id_cato=70001, id_afi='4444444',
        nro_solicitud='HR-2026-15', estado='REMITIDA_LEGAL',
        hruta_fecha=date(2026, 1, 20), cato_sup=1.25,
        cato_utm_xy='250100-8110200', cato_frec=2,
        renov_sup=1.0, renov_utm_xy='250150-8110250',
        vigencia_inicio=date(2026, 2, 1),
        fecha_vencimiento=date(2027, 2, 1),
        resol_nro='RA-33', resol_fecha=date(2026, 6, 1),
        resol_obs='Se aprueba la renovacion', resultado='APROBADA',
        tecnico_info_nro='CITE-T-5', legal_info_nro='CITE-L-9',
        federacion='FED-TEST', central='CENTRAL TEST',
        sindicato='SINDICATO TEST', departamento='COCHABAMBA',
        provincia='CHAPARE', municipio='VILLA TUNARI')
    db.session.add_all([cambio, renov])
    db.session.flush()

    informe = InformeVisitaTecnica(
        id_renovacion=renov.id, fecha_visita=date(2026, 3, 3),
        resultado='FACTIBLE', nro_informe='CITE-T-5',
        superficie=1.2, coordenadas='250100-8110200',
        causal_inciso='a', observaciones='Parcela verificada en campo',
        tecnico_nombre='TEC CAMPO', tecnico_ci='7070707')
    actuaciones = [
        ActuacionLegal(id_renovacion=renov.id, tipo='INFORME_LEGAL',
                       fecha=date(2026, 4, 1), nro_cite='CITE-L-9',
                       dictamen='PROCEDENTE',
                       contenido='Cumple los requisitos del DS 3318',
                       responsable_nombre='ABOG LEGAL',
                       responsable_cargo='ASESOR LEGAL'),
        ActuacionLegal(id_renovacion=renov.id, tipo='OBSERVACION_LEGAL',
                       fecha=date(2026, 4, 5), nro_cite='CITE-O-2',
                       contenido='Falta fotocopia de CI',
                       responsable_nombre='ABOG LEGAL',
                       responsable_cargo='ASESOR LEGAL'),
        ActuacionLegal(id_renovacion=renov.id, tipo='RESOLUCION',
                       fecha=date(2026, 6, 1), nro_cite='RA-33',
                       dictamen='APROBADA',
                       contenido='Se resuelve aprobar la renovacion',
                       responsable_nombre='ABOG LEGAL',
                       responsable_cargo='ASESOR LEGAL'),
    ]
    plano = Plano(nro_plano='RT-2026-001', id_cato=70001, id_afi='4444444',
                  tipo='MENSURA', estado='APROBADO',
                  fecha_registro=date(2026, 2, 15),
                  fecha_plano=date(2026, 2, 10), superficie=0.16,
                  coordenadas='250100-8110200', escala='1:500',
                  zona_utm='19S', dibujante='DIB TECNICO',
                  archivo_nombre='plano_rt2026001.pdf',
                  archivo_formato='PDF', ubicacion_fisica='ESTANTE-3')
    db.session.add_all([informe] + actuaciones + [plano])
    db.session.flush()

    revision = PlanoRevision(id_plano=plano.id_plano,
                             fecha_revision=date(2026, 2, 20),
                             resultado='APROBADO',
                             documentacion='Carpeta tecnica completa',
                             observaciones='Sin observaciones',
                             revisor_nombre='REV TECNICO')
    db.session.add(revision)
    db.session.commit()

    return {
        'id_trf': cambio.id_trf,
        'id_renov': renov.id,
        'id_informe': informe.id_informe,
        'id_informe_legal': actuaciones[0].id_actuacion,
        'id_observacion': actuaciones[1].id_actuacion,
        'id_resolucion': actuaciones[2].id_actuacion,
        'id_plano': plano.id_plano,
        'id_revision': revision.id_revision,
    }


# ---------------------------------------------------------------------------
# Ventanilla: afiliado, cato, cambio
# ---------------------------------------------------------------------------

def test_doc_afiliado(client, auth_admin):
    texto = _texto_docx(client.get('/api/documentos/afiliado/4444444',
                                   headers=auth_admin))
    assert '4444444' in texto
    assert 'PEREZ MAMANI JUAN' in texto
    assert 'SINDICATO TEST' in texto     # relacion cato -> org sindical
    assert '70001' in texto              # cato del afiliado


def test_doc_afiliado_inexistente(client, auth_admin):
    resp = client.get('/api/documentos/afiliado/0000000',
                      headers=auth_admin)
    assert resp.status_code == 404


def test_doc_cato(client, auth_admin):
    texto = _texto_docx(client.get('/api/documentos/cato/70001',
                                   headers=auth_admin))
    assert '70001' in texto
    assert 'PEREZ MAMANI JUAN' in texto  # relacion cato -> afiliado
    assert 'FED-TEST' in texto           # jerarquia sindical completa


def test_doc_cato_inexistente(client, auth_admin):
    assert client.get('/api/documentos/cato/99999',
                      headers=auth_admin).status_code == 404


def test_doc_cambio(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/cambio/{datos_documentos['id_trf']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'COMPRA-VENTA' in texto
    assert 'PEREZ MAMANI JUAN' in texto  # titular anterior
    assert 'QUISPE ROJAS MARIA' in texto  # nuevo titular
    assert '10/05/2026' in texto         # fecha formateada dd/mm/aaaa


def test_doc_cambio_inexistente(client, auth_admin):
    assert client.get('/api/documentos/cambio/99999',
                      headers=auth_admin).status_code == 404


# ---------------------------------------------------------------------------
# Renovaciones: solicitud, informe tecnico, resolucion
# ---------------------------------------------------------------------------

def test_doc_renovacion_solicitud(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/renovacion/{datos_documentos['id_renov']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'HR-2026-15' in texto
    assert 'PEREZ MAMANI JUAN' in texto
    assert 'SINDICATO TEST' in texto
    assert '250100-8110200' in texto     # coordenadas parcela anterior


def test_doc_renovacion_informe(client, auth_admin, datos_documentos):
    ruta = (f"/api/documentos/renovacion/{datos_documentos['id_renov']}"
            '?tipo=informe')
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'CITE-T-5' in texto
    assert 'FACTIBLE' in texto
    assert 'TEC CAMPO' in texto
    assert 'Parcela verificada en campo' in texto


def test_doc_renovacion_informe_especifico(client, auth_admin,
                                           datos_documentos):
    ruta = (f"/api/documentos/renovacion/{datos_documentos['id_renov']}"
            f"?tipo=informe&id_informe={datos_documentos['id_informe']}")
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'CITE-T-5' in texto


def test_doc_renovacion_resolucion(client, auth_admin, datos_documentos):
    ruta = (f"/api/documentos/renovacion/{datos_documentos['id_renov']}"
            '?tipo=resolucion')
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'RA-33' in texto
    assert 'APROBADA' in texto
    assert '01/02/2027' in texto         # fecha de vencimiento
    assert 'ABOG LEGAL' in texto         # responsable de la resolucion


def test_doc_renovacion_tipo_invalido(client, auth_admin, datos_documentos):
    ruta = (f"/api/documentos/renovacion/{datos_documentos['id_renov']}"
            '?tipo=nada')
    assert client.get(ruta, headers=auth_admin).status_code == 400


def test_doc_renovacion_inexistente(client, auth_admin):
    assert client.get('/api/documentos/renovacion/99999',
                      headers=auth_admin).status_code == 404


def test_doc_renovacion_sin_informes(client, auth_admin, datos_documentos):
    renov = RenovacionProgramada(id_renov=901, id_cato=70001,
                                 id_afi='4444444', estado='PENDIENTE')
    db.session.add(renov)
    db.session.commit()
    resp = client.get(f'/api/documentos/renovacion/{renov.id}?tipo=informe',
                      headers=auth_admin)
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Legal: informe, observacion, resolucion (segun tipo de la actuacion)
# ---------------------------------------------------------------------------

def test_doc_legal_informe(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/legal/{datos_documentos['id_informe_legal']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'CITE-L-9' in texto
    assert 'PROCEDENTE' in texto
    assert 'Cumple los requisitos del DS 3318' in texto
    assert 'PEREZ MAMANI JUAN' in texto  # relacion actuacion -> afiliado


def test_doc_legal_observacion(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/legal/{datos_documentos['id_observacion']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'CITE-O-2' in texto
    assert 'Falta fotocopia de CI' in texto


def test_doc_legal_resolucion(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/legal/{datos_documentos['id_resolucion']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'RA-33' in texto
    assert 'Se resuelve aprobar la renovacion' in texto


def test_doc_legal_inexistente(client, auth_admin):
    assert client.get('/api/documentos/legal/99999',
                      headers=auth_admin).status_code == 404


# ---------------------------------------------------------------------------
# Planos: certificado y acta de revision
# ---------------------------------------------------------------------------

def test_doc_certificado_plano(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/plano/{datos_documentos['id_plano']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'RT-2026-001' in texto
    assert 'MENSURA' in texto
    assert 'DIB TECNICO' in texto
    assert 'PEREZ MAMANI JUAN' in texto
    assert '0.16' in texto               # superficie sin ceros sobrantes


def test_doc_plano_inexistente(client, auth_admin):
    assert client.get('/api/documentos/plano/99999',
                      headers=auth_admin).status_code == 404


def test_doc_acta_revision(client, auth_admin, datos_documentos):
    ruta = f"/api/documentos/plano-revision/{datos_documentos['id_revision']}"
    texto = _texto_docx(client.get(ruta, headers=auth_admin))
    assert 'RT-2026-001' in texto        # relacion revision -> plano
    assert 'APROBADO' in texto
    assert 'Carpeta tecnica completa' in texto
    assert 'REV TECNICO' in texto


def test_doc_revision_inexistente(client, auth_admin):
    assert client.get('/api/documentos/plano-revision/99999',
                      headers=auth_admin).status_code == 404


# ---------------------------------------------------------------------------
# Seguridad
# ---------------------------------------------------------------------------

def test_doc_requiere_token(client, datos_documentos):
    assert client.get('/api/documentos/afiliado/4444444').status_code == 401
    ruta = f"/api/documentos/plano/{datos_documentos['id_plano']}"
    assert client.get(ruta).status_code == 401


def test_doc_lectura_para_todos_los_roles(client, auth_lector,
                                          datos_documentos):
    """USR_PLANOS (solo lectura) tambien puede descargar documentos."""
    resp = client.get('/api/documentos/afiliado/4444444',
                      headers=auth_lector)
    assert resp.status_code == 200
